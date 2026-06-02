import os
import cv2
import imagehash
import numpy as np

from PIL import Image
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


VIDEO_PATH = "video.mp4"

OUTPUT_DIR = Path("screenshots")

DOCX_FILE = "output.docx"

OUTPUT_DIR.mkdir(exist_ok=True)


# ==========================
# 参数
# ==========================

# 每秒分析多少帧
SAMPLE_FPS = 5

# 画面变化阈值
CHANGE_THRESHOLD = 3

# 连续稳定秒数
STABLE_SECONDS = 2

# 哈希距离阈值
HASH_THRESHOLD = 5


# ==========================
# 提取稳定页面
# ==========================

cap = cv2.VideoCapture(VIDEO_PATH)

fps = cap.get(cv2.CAP_PROP_FPS)

frame_interval = int(fps / SAMPLE_FPS)

stable_needed = SAMPLE_FPS * STABLE_SECONDS

ret, prev_frame = cap.read()

if not ret:
    raise Exception("无法读取视频")

stable_count = 0
saved_count = 0

last_hash = None

frame_index = 0

print("开始分析视频...")

while True:

    ret, frame = cap.read()

    if not ret:
        break

    frame_index += 1

    if frame_index % frame_interval != 0:
        continue

    gray1 = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)

    gray2 = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

    diff = cv2.absdiff(gray1, gray2)

    score = np.mean(diff)

    if score < CHANGE_THRESHOLD:

        stable_count += 1

    else:

        stable_count = 0

    if stable_count >= stable_needed:

        temp_file = OUTPUT_DIR / f"temp_{saved_count}.png"

        cv2.imwrite(str(temp_file), frame)

        img = Image.open(temp_file)

        current_hash = imagehash.phash(img)

        save_this = False

        if last_hash is None:
            save_this = True

        else:
            distance = current_hash - last_hash

            if distance > HASH_THRESHOLD:
                save_this = True

        if save_this:

            final_file = OUTPUT_DIR / f"page_{saved_count:03d}.png"

            os.rename(temp_file, final_file)

            last_hash = current_hash

            print("保存:", final_file)

            saved_count += 1

        else:

            os.remove(temp_file)

        stable_count = 0

    prev_frame = frame.copy()

cap.release()

print(f"提取完成，共 {saved_count} 张截图")


# ==========================
# 生成Word
# ==========================

print("生成Word...")

COLS = 3            # 每行 3 张
ROWS_PER_PAGE = 2   # 每页 2 行，共 6 张
IMAGES_PER_PAGE = COLS * ROWS_PER_PAGE
IMAGE_WIDTH_CM = 4.5  # A4 可用宽约 14.6cm，3 列各 4.5cm


def remove_table_borders(table):
    """去除表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=0, left=30, bottom=0, right=30):
    """设置单元格内边距（单位：twip，1cm≈567twip）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def zero_paragraph_spacing(paragraph):
    """清除段落前后间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Cm(0)
    pf.space_after = Cm(0)


doc = Document()

doc.add_heading("视频截图整理", level=1)

image_list = sorted(OUTPUT_DIR.glob("page_*.png"))

for page_idx in range(0, len(image_list), IMAGES_PER_PAGE):

    if page_idx > 0:
        doc.add_page_break()

    page_images = image_list[page_idx: page_idx + IMAGES_PER_PAGE]
    rows_needed = (len(page_images) + COLS - 1) // COLS

    table = doc.add_table(rows=rows_needed, cols=COLS)
    remove_table_borders(table)

    for i, img_path in enumerate(page_images):
        row_idx = i // COLS
        col_idx = i % COLS
        cell = table.cell(row_idx, col_idx)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        zero_paragraph_spacing(paragraph)
        run = paragraph.add_run()
        run.add_picture(str(img_path), width=Cm(IMAGE_WIDTH_CM))

doc.save(DOCX_FILE)

print("完成:", DOCX_FILE)